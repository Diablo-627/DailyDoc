import os
import re
import asyncio
import logging
import shutil
import tempfile
from aiogram import Bot, types, F, Router
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import (
    Message,
    FSInputFile,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    ReplyKeyboardRemove
)
from docx import Document
from docx.shared import Cm
from docx.oxml import parse_xml
from PIL import Image
import xml.etree.ElementTree as ET
import zipfile

logger = logging.getLogger(__name__)
garbage_router = Router()

# Константы
PHOTO_WIDTH = 13.33  # см
PHOTO_HEIGHT = 7.5   # см
CM_TO_PX = 37.8      # 1 см ≈ 37.8 пикселей
TEMPLATE_NAME = "template21.docx"  # Название файла шаблона

class GarbageReportState(StatesGroup):
    DATE = State()
    ADDRESSES = State()
    EQUIPMENT = State()
    GARBAGE_AMOUNT = State()
    PARTICIPANTS = State()
    HOURS = State()
    INPUT_PHOTOS = State()
    PHOTO_ASSIGNMENT = State()

async def start_garbage_report(message: types.Message, state: FSMContext):
    """Запуск сценария отчета по вывозу мусора"""
    await state.clear()
    await state.set_state(GarbageReportState.DATE)
    await message.answer(
        "📅 Введите дату вывоза мусора:",
        reply_markup=ReplyKeyboardRemove()
    )

@garbage_router.message(GarbageReportState.DATE)
async def process_date(message: types.Message, state: FSMContext):
    """Обработка даты - принимаем любой текст"""
    await state.update_data(date=message.text)
    await state.set_state(GarbageReportState.ADDRESSES)
    await message.answer(
        "🏠 Введите адреса (каждый адрес с новой строки):\n"
        "Пример:\n"
        "Ул. Ленина, д. 1\n"
        "Пр. Мира, д. 15\n"
        "Ул. Центральная, д. 8"
    )

@garbage_router.message(GarbageReportState.ADDRESSES)
async def process_addresses(message: types.Message, state: FSMContext):
    """Обработка списка адресов"""
    addresses = [addr.strip() for addr in message.text.split('\n') if addr.strip()]
    if not addresses:
        await message.answer("❌ Адреса не введены. Введите хотя бы один адрес")
        return
        
    await state.update_data(
        addresses=addresses,
        photos={addr: [] for addr in addresses},
        photo_counter=0
    )
    await state.set_state(GarbageReportState.EQUIPMENT)
    await message.answer("🚛 Введите задействованную технику:")

@garbage_router.message(GarbageReportState.EQUIPMENT)
async def process_equipment(message: types.Message, state: FSMContext):
    """Обработка информации о технике - принимаем любой текст"""
    await state.update_data(equipment=message.text)
    await state.set_state(GarbageReportState.GARBAGE_AMOUNT)
    await message.answer("🗑️ Введите количество вывезенного мусора (в тоннах):")

@garbage_router.message(GarbageReportState.GARBAGE_AMOUNT)
async def process_garbage_amount(message: types.Message, state: FSMContext):
    """Обработка объема мусора - принимаем любой текст"""
    await state.update_data(garbage_amount=message.text)
    await state.set_state(GarbageReportState.PARTICIPANTS)
    await message.answer("👥 Введите количество участников:")

@garbage_router.message(GarbageReportState.PARTICIPANTS)
async def process_participants(message: types.Message, state: FSMContext):
    """Обработка количества участников - принимаем любой текст"""
    await state.update_data(participants=message.text)
    await state.set_state(GarbageReportState.HOURS)
    await message.answer("⏱️ Введите количество часов работы техники:")

@garbage_router.message(GarbageReportState.HOURS)
async def process_hours(message: types.Message, state: FSMContext):
    """Обработка информации о часах работы - принимаем любой текст"""
    await state.update_data(hours=message.text)
    data = await state.get_data()
    
    total_photos = len(data['addresses']) * 2
    await state.set_state(GarbageReportState.INPUT_PHOTOS)
    await message.answer(
        f"📸 Теперь загрузите {total_photos} фото (по 2 на каждый адрес).\n"
        f"Порядок адресов:\n" + "\n".join(
            f"{i+1}. {addr}" for i, addr in enumerate(data['addresses'])
        )
    )

# Обработчик для не-фото в состоянии INPUT_PHOTOS
@garbage_router.message(GarbageReportState.INPUT_PHOTOS, F.photo)
async def process_photo_upload(message: types.Message, state: FSMContext):
    """Обработка загруженных фото (поддержка нескольких фото за раз)"""
    data = await state.get_data()
    addresses = data['addresses']
    total_photos = len(addresses) * 2
    
    # Получаем file_id самого большого варианта фото (последний в списке)
    photo_file_ids = [photo.file_id for photo in message.photo]
    largest_photo_id = photo_file_ids[-1]  # Берем последний (самый большой вариант)
    
    # Если несколько фото в одном сообщении (альбом)
    if len(photo_file_ids) > 1:
        await state.update_data(
            photo_buffer=photo_file_ids,
            processing_album=True
        )
        await message.answer("📸 Получено несколько фото. Начнем их распределение...")
        await process_next_photo_from_buffer(message, state)
        return
    
    # Обработка одного фото
    await state.update_data(
        current_photo=largest_photo_id,
        photo_counter=data.get('photo_counter', 0) + 1
    )
    await ask_photo_assignment(message, state)

async def process_next_photo_from_buffer(message: types.Message, state: FSMContext):
    """Обработка следующего фото из буфера"""
    data = await state.get_data()
    photo_buffer = data.get('photo_buffer', [])
    
    if not photo_buffer:
        await state.update_data(processing_album=False)
        await check_completion(message, state)
        return
    
    current_photo = photo_buffer.pop(0)
    await state.update_data(
        current_photo=current_photo,
        photo_buffer=photo_buffer,
        photo_counter=data.get('photo_counter', 0) + 1
    )
    await ask_photo_assignment(message, state)

async def ask_photo_assignment(message: types.Message, state: FSMContext):
    """Запрос привязки фото к адресу"""
    data = await state.get_data()
    addresses = data['addresses']
    photo_counter = data['photo_counter']
    total_photos = len(addresses) * 2
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[])
    for address in addresses:
        photo_count = len(data['photos'].get(address, []))
        if photo_count < 2:
            keyboard.inline_keyboard.append([
                InlineKeyboardButton(
                    text=f"{address} ({photo_count+1}/2)", 
                    callback_data=f"address_{address}"
                )
            ])
    
    if not keyboard.inline_keyboard:
        await message.answer("❌ Все фото уже распределены!")
        await check_completion(message, state)
        return
        
    await state.set_state(GarbageReportState.PHOTO_ASSIGNMENT)
    await message.answer(
        f"📎 Фото {photo_counter}/{total_photos}\n"
        "Выберите адрес для этого фото:",
        reply_markup=keyboard
    )

async def check_completion(message: types.Message, state: FSMContext):
    """Проверка завершения загрузки фото"""
    data = await state.get_data()
    total_photos = len(data['addresses']) * 2
    
    if data.get('photo_counter', 0) >= total_photos:
        await generate_garbage_report(message, state)
    else:
        await state.set_state(GarbageReportState.INPUT_PHOTOS)
        remaining = total_photos - data.get('photo_counter', 0)
        await message.answer(
            f"📸 Осталось загрузить {remaining} фото. "
            "Можете отправить одно или несколько фото сразу."
        )

@garbage_router.callback_query(GarbageReportState.PHOTO_ASSIGNMENT, F.data.startswith("address_"))
async def assign_photo_to_address(callback: types.CallbackQuery, state: FSMContext):
    """Привязка фото к адресу"""
    address = callback.data.split("_", 1)[1]
    data = await state.get_data()
    current_photo = data['current_photo']
    
    photos = data['photos'].copy()
    if address in photos and len(photos[address]) < 2:
        photos[address].append(current_photo)
    else:
        await callback.answer("❌ Нельзя добавить больше 2 фото на адрес!")
        return
        
    photo_counter = data['photo_counter']
    total_photos = len(data['addresses']) * 2
    
    await state.update_data(photos=photos)
    await callback.answer(f"✅ Фото привязано к адресу: {address}")
    
    try:
        await callback.message.delete()
    except Exception as e:
        logger.error(f"Ошибка при удалении сообщения: {e}")
    
    if photo_counter >= total_photos:
        await generate_garbage_report(callback.message, state)
        return
        
    await state.set_state(GarbageReportState.INPUT_PHOTOS)
    await callback.message.answer(
        f"✅ Фото {photo_counter}/{total_photos} сохранено\n"
        "Загружайте следующее фото:"
    )

def apply_photo_style(run):
    """Применение стилей к фото в документе"""
    effect = parse_xml(
        r'<a:prstGeom prst="roundRect" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"/>'
    )
    run._element.xpath('.//pic:spPr')[0].append(effect)
    
    ln = parse_xml(
        r'<a:ln w="12700" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        r'<a:solidFill><a:srgbClr val="FFFFFF"/></a:solidFill>'
        r'</a:ln>'
    )
    run._element.xpath('.//pic:spPr')[0].append(ln)

async def download_and_process_photo(file_id: str, bot: Bot, target_width: float, target_height: float):
    """Скачивание и обработка фото"""
    file = await bot.get_file(file_id)
    photo_data = await bot.download_file(file.file_path)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp_input:
        # Сохраняем оригинальное фото
        temp_input.write(photo_data.read())
        temp_input.flush()
        
        with Image.open(temp_input.name) as img:
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            target_width_px = int(target_width * CM_TO_PX)
            target_height_px = int(target_height * CM_TO_PX)
            
            # Рассчитываем масштаб
            width, height = img.size
            scale = max(
                target_width_px / width,
                target_height_px / height
            )
            scaled_width = int(width * scale)
            scaled_height = int(height * scale)
            
            # Масштабируем и обрезаем
            img = img.resize((scaled_width, scaled_height), Image.LANCZOS)
            left = (scaled_width - target_width_px) // 2
            top = (scaled_height - target_height_px) // 2
            img = img.crop((
                max(0, left),
                max(0, top),
                min(scaled_width, left + target_width_px),
                min(scaled_height, top + target_height_px)
            ))
            
            temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
            img.save(temp_output, format='JPEG', quality=95)
            temp_output.close()
            
            os.unlink(temp_input.name)
            return temp_output.name

def replace_text_in_docx(doc_path: str, replacements: dict):
    """Замена текста в docx файле"""
    with tempfile.TemporaryDirectory() as tmp_dir:
        with zipfile.ZipFile(doc_path, 'r') as zip_ref:
            zip_ref.extractall(tmp_dir)
        
        document_xml_path = os.path.join(tmp_dir, 'word', 'document.xml')
        tree = ET.parse(document_xml_path)
        root = tree.getroot()
        
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        ET.register_namespace('w', namespaces['w'])
        
        for elem in root.iter():
            if elem.text:
                for placeholder, value in replacements.items():
                    if placeholder in elem.text:
                        elem.text = elem.text.replace(placeholder, value)
        
        tree.write(document_xml_path, encoding='UTF-8', xml_declaration=True)
        
        with zipfile.ZipFile(doc_path, 'w') as zip_ref:
            for root_dir, _, files in os.walk(tmp_dir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, tmp_dir)
                    zip_ref.write(file_path, arcname)

async def generate_garbage_report(message: types.Message, state: FSMContext):
    data = await state.get_data()
    bot = Bot.get_current()
    
    with tempfile.TemporaryDirectory() as temp_dir:
        doc_path = os.path.join(temp_dir, "Отчет_вывоза_мусора.docx")
        template_path = TEMPLATE_NAME
        shutil.copy(template_path, doc_path)
        
        # Основные замены
        replacements = {
            "<<DATE>>": data.get('date', ''),
            "<<EQUIPMENT>>": data.get('equipment', ''),
            "<<GARBAGE_AMOUNT>>": data.get('garbage_amount', ''),
            "<<PARTICIPANTS>>": data.get('participants', ''),
            "<<HOURS>>": data.get('hours', ''),
            "<<ADDRESSES>>": "\n".join(data['addresses'])
        }
        
        # Замена текста
        replace_text_in_docx(doc_path, replacements)
        
        # Загрузка документа
        doc = Document(doc_path)
        
        # Находим секцию для дублирования
        section_start = None
        section_end = None
        for i, para in enumerate(doc.paragraphs):
            if "<<ADDRESS_SECTION>>" in para.text:
                section_start = i
            if "<<END_SECTION>>" in para.text and section_start is not None:
                section_end = i
                break
        
        if section_start is None or section_end is None:
            await message.answer("❌ В шаблоне не найдены маркеры секций!")
            await state.clear()
            return
        
        # Удаляем маркеры секций
        del doc.paragraphs[section_end]
        del doc.paragraphs[section_start]
        
        # Копируем и вставляем секции для каждого адреса
        template_paragraphs = list(doc.paragraphs[section_start:section_end])
        
        for i, address in enumerate(data['addresses']):
            # Для первого адреса используем существующую секцию
            if i > 0:
                # Добавляем разрыв страницы перед новой секцией
                doc.add_page_break()
                
                # Копируем шаблонные параграфы
                for para in template_paragraphs:
                    new_para = doc.add_paragraph()
                    new_para.text = para.text
                    new_para.style = para.style
            
            # Замена в текущей секции
            start_idx = section_start if i == 0 else len(doc.paragraphs) - len(template_paragraphs)
            
            for j in range(start_idx, start_idx + len(template_paragraphs)):
                para = doc.paragraphs[j]
                if "<<CURRENT_ADDRESS>>" in para.text:
                    para.text = para.text.replace("<<CURRENT_ADDRESS>>", address)
                
                # Вставка фото
                for k in range(1, 3):
                    placeholder = f"<<PHOTO_{k}>>"
                    if placeholder in para.text:
                        if k <= len(data['photos'][address]):
                            try:
                                photo_path = await download_and_process_photo(
                                    data['photos'][address][k-1], 
                                    bot,
                                    PHOTO_WIDTH,
                                    PHOTO_HEIGHT
                                )
                                para.text = para.text.replace(placeholder, "")
                                run = para.add_run()
                                run.add_picture(photo_path, width=Cm(PHOTO_WIDTH), height=Cm(PHOTO_HEIGHT))
                                apply_photo_style(run)
                                os.unlink(photo_path)
                            except Exception as e:
                                logger.error(f"Ошибка обработки фото: {e}")
                                para.text = para.text.replace(placeholder, f"❌ Ошибка загрузки фото {k}")
        
        doc.save(doc_path)
        await message.answer("✅ Отчет готов!")
        await message.answer_document(FSInputFile(doc_path, filename="Отчет_вывоза_мусора.docx"))
    
        await state.clear()
        
        # Загрузка документа для вставки фото
        doc = Document(doc_path)
        
        # Обработка и вставка фотографий
        for i, address in enumerate(addresses):
            for j in range(2):
                placeholder = f"<<PHOTO_{i+1}_{j+1}>>"
                if j < len(data['photos'][address]):
                    file_id = data['photos'][address][j]
                    try:
                        photo_path = await download_and_process_photo(
                            file_id, bot, PHOTO_WIDTH, PHOTO_HEIGHT
                        )
                        
                        found = False
                        for para in doc.paragraphs:
                            if placeholder in para.text:
                                para.text = para.text.replace(placeholder, '')
                                run = para.add_run()
                                run.add_picture(photo_path, width=Cm(PHOTO_WIDTH), height=Cm(PHOTO_HEIGHT))
                                apply_photo_style(run)
                                found = True
                                break
                        
                        if not found:
                            logger.warning(f"Не найден плейсхолдер для фото: {placeholder}")
                        
                        os.unlink(photo_path)
                    except Exception as e:
                        logger.error(f"Ошибка обработки фото: {e}")
                        await message.answer(f"❌ Ошибка обработки фото для адреса {address}")
        
        doc.save(doc_path)
        await message.answer("✅ Отчет готов!")
        await message.answer_document(FSInputFile(doc_path, filename="Отчет_вывоза_мусора.docx"))
    
    await state.clear()

# Экспорт для использования в других модулях
__all__ = ['garbage_router', 'start_garbage_report']
