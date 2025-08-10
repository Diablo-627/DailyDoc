import os
import asyncio
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
from aiogram import Router
from aiogram.types import Message, CallbackQuery
from aiogram.filters import Command
from aiogram.utils.keyboard import InlineKeyboardMarkup, InlineKeyboardButton
from docx import Document
from PIL import Image

MAX_PHOTOS = 15
PHOTO_SIZES = {
    "default": (800, 600)
}

class DailyReport:
    def __init__(self, bot, template_path: str, photos_dir: str, temp_dir: str):
        self.bot = bot
        self.template_path = template_path
        self.photos_dir = photos_dir
        self.temp_dir = temp_dir
        os.makedirs(self.photos_dir, exist_ok=True)
        os.makedirs(self.temp_dir, exist_ok=True)

        self.router = Router()
        self.sessions = {}
        self.executor = ThreadPoolExecutor(max_workers=4)

        self.router.message(Command("reset"))(self._reset_from_route)
        self.router.message(Command("help"))(self._help_from_route)
        self.router.message(lambda m: m.photo)(self._handle_photo_only)
        self.router.message(lambda m: m.text and m.text.strip())(self._handle_text)
        self.router.callback_query(lambda c: c.data.startswith("tag_"))(self._handle_photo_tag)

        self.fields_order = [
            "Название объекта", "Адрес", "Ответственный", "Дата"
        ]
        self.photo_tags = [f"PHOTO_{i}" for i in range(1, MAX_PHOTOS + 1)]

    def _get_or_create_session(self, chat_id: int):
        if chat_id not in self.sessions:
            self.sessions[chat_id] = {
                "fields": {field: "" for field in self.fields_order},
                "photos": {},
                "photo_queue": [],
                "remaining_tags": self.photo_tags.copy(),
                "current_file_id": None,
                "processing": False,
                "state": "text_input",
                "current_field_index": 0
            }
        return self.sessions[chat_id]

    async def start_for_user(self, chat_id: int):
        session = self._get_or_create_session(chat_id)
        session["state"] = "text_input"
        session["current_field_index"] = 0
        await self.bot.send_message(chat_id, f"Введите: {self.fields_order[0]}")

    async def _reset_from_route(self, message: Message):
        self.sessions.pop(message.chat.id, None)
        await message.answer("Сессия сброшена. Начните заново через /start.")

    async def _help_from_route(self, message: Message):
        await message.answer("Ежедневный отчёт. Сначала вводите текстовые поля, потом отправляете фото (до 15 штук).")

    async def _handle_text(self, message: Message):
        chat_id = message.chat.id
        session = self._get_or_create_session(chat_id)

        if session["state"] != "text_input":
            return

        field_name = self.fields_order[session["current_field_index"]]
        session["fields"][field_name] = message.text.strip()
        session["current_field_index"] += 1

        if session["current_field_index"] < len(self.fields_order):
            next_field = self.fields_order[session["current_field_index"]]
            await message.answer(f"Введите: {next_field}")
        else:
            session["state"] = "photo_input"
            await message.answer("Теперь отправьте фото (до 15 штук).")

    async def _handle_photo_only(self, message: Message):
        chat_id = message.chat.id
        session = self._get_or_create_session(chat_id)

        if session["state"] != "photo_input":
            return

        if len(session["photos"]) >= MAX_PHOTOS or not session["remaining_tags"]:
            await message.answer("Все фото уже получены.")
            return

        session["photo_queue"].append(message.photo[-1].file_id)
        if not session["processing"]:
            await self._process_next_photo(chat_id)

    async def _process_next_photo(self, chat_id: int):
        session = self._get_or_create_session(chat_id)
        if not session["photo_queue"]:
            return
        session["processing"] = True
        file_id = session["photo_queue"][0]
        session["current_file_id"] = file_id

        kb = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text=tag, callback_data=f"tag_{tag}")] for tag in session["remaining_tags"]
        ] + [[InlineKeyboardButton(text="Пропустить", callback_data="tag_skip")]])

        await self.bot.send_message(chat_id, "Выберите тег для фото:", reply_markup=kb)

    async def _handle_photo_tag(self, callback: CallbackQuery):
        chat_id = callback.message.chat.id
        session = self._get_or_create_session(chat_id)
        tag = callback.data.replace("tag_", "")

        try:
            await callback.message.delete()
        except:
            pass

        if tag == "skip":
            session["photo_queue"].pop(0)
            session["current_file_id"] = None
            session["processing"] = False
            if session["photo_queue"]:
                await self._process_next_photo(chat_id)
            return

        photo_path = os.path.join(self.photos_dir, f"{chat_id}_{tag}.jpg")
        file = await self.bot.get_file(session["current_file_id"])
        await self.bot.download_file(file.file_path, photo_path)

        width, height = PHOTO_SIZES.get(tag, PHOTO_SIZES["default"])
        loop = asyncio.get_running_loop()
        await loop.run_in_executor(self.executor, self._resize_and_crop_image, photo_path, width, height)

        session["photos"][tag] = photo_path
        if tag in session["remaining_tags"]:
            session["remaining_tags"].remove(tag)

        session["photo_queue"].pop(0)
        session["current_file_id"] = None
        session["processing"] = False

        if len(session["photos"]) >= MAX_PHOTOS or not session["remaining_tags"]:
            await self._generate_docx(chat_id)
            return

        if session["photo_queue"]:
            await self._process_next_photo(chat_id)

    def _resize_and_crop_image(self, path, width, height):
        img = Image.open(path)
        img = img.resize((width, height), Image.LANCZOS)
        img.save(path)

    async def _generate_docx(self, chat_id: int):
        session = self.sessions.get(chat_id)
        if not session:
            return

        doc = Document(self.template_path)
        for p in doc.paragraphs:
            for field, value in session["fields"].items():
                if f"{{{field}}}" in p.text:
                    p.text = p.text.replace(f"{{{field}}}", value)

        for tag, path in session["photos"].items():
            for p in doc.paragraphs:
                if f"{{{tag}}}" in p.text:
                    p.text = ""
                    run = p.add_run()
                    run.add_picture(path)

        output_path = os.path.join(self.temp_dir, f"report_{chat_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(output_path)

        await self.bot.send_document(chat_id, open(output_path, "rb"))
        await self.bot.send_message(chat_id, "Отчёт готов.")

        self._cleanup_files(session)
        session["state"] = None

    def _cleanup_files(self, session):
        for path in session["photos"].values():
            try:
                os.remove(path)
            except:
                pass
        session["photos"].clear()
